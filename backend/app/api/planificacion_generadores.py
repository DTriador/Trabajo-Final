# backend/app/api/planificacion_generadores.py
"""
Generación de los documentos descargables de una planificación (.docx / .pdf).

_generar_docx_planificacion: ya existía como función propia en
router_planificacion.py, se movió tal cual (con el fix de anchos de
columna + formateo de fecha ya aplicado).

_generar_pdf_planificacion: antes era código pegado directamente dentro
del endpoint /planificacion/{id_plan}/exportar-pdf. Se extrajo a una
función propia (misma lógica, sin cambios) para que el router quede
liviano y el endpoint solo arme la respuesta HTTP.
"""
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

from app.api.planificacion_helpers import _formatear_fecha_clase


# ── Word (.docx) ─────────────────────────────────────────────────────────────

def _generar_docx_planificacion(plan: dict, clases: list) -> bytes:
    """
    Construye el .docx de una planificación (título + datos generales + cronograma).
    Reutilizado tanto por el endpoint de descarga bajo demanda (/exportar-word)
    como por el guardado automático en Mis Materiales al crear la planificación
    desde el wizard.
    """
    doc = Document()

    # Título
    titulo = doc.add_heading(plan.get("nombre_clase", "Planificación"), 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Datos generales ──────────────────────────────────────────────
    doc.add_heading("Datos generales", level=1)
    tabla_meta = doc.add_table(rows=3, cols=2)
    tabla_meta.style = "Table Grid"
    tabla_meta.autofit = False

    # Anchos fijos: etiqueta angosta, valor ancho — suman ~6.3" (dentro de márgenes A4/carta)
    ancho_label = Inches(1.6)
    ancho_valor = Inches(4.7)
    for row in tabla_meta.rows:
        row.cells[0].width = ancho_label
        row.cells[1].width = ancho_valor

    meta = [
        ("Contenido mínimo",  plan.get("contenido_minimo", "—")),
        ("Duración de clase", plan.get("duracion", "—")),
        ("Total de clases",   str(len([c for c in clases if c.get("tipo") == "clase"]))),
    ]
    for i, (k, v) in enumerate(meta):
        celda_label = tabla_meta.rows[i].cells[0]
        celda_label.text = k
        celda_label.paragraphs[0].runs[0].bold = True

        celda_valor = tabla_meta.rows[i].cells[1]
        celda_valor.text = v or "—"
        # Tamaño de fuente un poco menor para que el texto largo entre prolijo
        for p in celda_valor.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Cronograma ────────────────────────────────────────────────────
    doc.add_heading("Cronograma", level=1)
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"
    tabla.autofit = False

    # Anchos: N° angosto, Fecha media, Tipo angosto, Tema el más ancho
    anchos = [Inches(0.4), Inches(1.7), Inches(1.1), Inches(3.1)]
    for i, enc in enumerate(["N°", "Fecha", "Tipo", "Tema"]):
        cell = tabla.rows[0].cells[i]
        cell.text = enc
        cell.paragraphs[0].runs[0].bold = True
        cell.width = anchos[i]

    tipo_labels = {"clase": "Clase", "examen": "Examen", "recuperatorio": "Recuperatorio"}
    for c in clases:
        row = tabla.add_row()
        row.cells[0].text = str(c.get("numero", ""))
        row.cells[1].text = _formatear_fecha_clase(c.get("fecha_programada", ""))
        row.cells[2].text = tipo_labels.get(c.get("tipo", "clase"), c.get("tipo", ""))
        row.cells[3].text = c.get("tema_clase", "")
        for i, cell in enumerate(row.cells):
            cell.width = anchos[i]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF ────────────────────────────────────────────────────────────────────

def _generar_pdf_planificacion(plan: dict, clases: list) -> bytes:
    """
    Construye el .pdf de una planificación (título + datos generales + cronograma).
    Extraído tal cual desde el endpoint /planificacion/{id_plan}/exportar-pdf
    para que ese endpoint quede liviano.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=50, bottomMargin=40)
    styles = getSampleStyleSheet()
    elementos = []

    # Título
    elementos.append(Paragraph(plan.get("nombre_clase", "Planificación"), styles["Title"]))
    elementos.append(Spacer(1, 12))

    # Datos generales
    elementos.append(Paragraph("Datos generales", styles["Heading2"]))
    meta_data = [
        ["Contenido mínimo", plan.get("contenido_minimo", "—") or "—"],
        ["Duración de clase",  plan.get("duracion", "—") or "—"],
        ["Total de clases",    str(len([c for c in clases if c.get("tipo") == "clase"]))],
    ]
    t_meta = Table(meta_data, colWidths=[140, 360])
    t_meta.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (0, -1), colors.HexColor("#e0f2fe")),
        ("FONTNAME",    (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE",    (0, 0), (-1, -1), 9),
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("PADDING",     (0, 0), (-1, -1), 6),
    ]))
    elementos.append(t_meta)
    elementos.append(Spacer(1, 16))

    # Cronograma
    elementos.append(Paragraph("Cronograma de clases", styles["Heading2"]))
    tipo_labels = {"clase": "Clase", "examen": "Examen", "recuperatorio": "Recuperatorio"}
    tipo_colores = {"clase": "#dbeafe", "examen": "#fef3c7", "recuperatorio": "#dcfce7"}
    tabla_data = [["N°", "Fecha", "Tipo", "Tema"]]
    for c in clases:
        tabla_data.append([
            str(c.get("numero", "")),
            _formatear_fecha_clase(c.get("fecha_programada", "")),
            tipo_labels.get(c.get("tipo", ""), c.get("tipo", "")),
            Paragraph(c.get("tema_clase", ""), styles["Normal"]),
        ])
    t = Table(tabla_data, colWidths=[25, 75, 80, 320])
    style_cmds = [
        ("BACKGROUND",  (0, 0), (-1, 0),  colors.HexColor("#1e3a8a")),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 8),
        ("GRID",        (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
        ("PADDING",     (0, 0), (-1, -1), 5),
    ]
    for i, c in enumerate(clases, start=1):
        bg = tipo_colores.get(c.get("tipo", "clase"), "#ffffff")
        style_cmds.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor(bg)))
    t.setStyle(TableStyle(style_cmds))
    elementos.append(t)

    doc.build(elementos)
    buf.seek(0)
    return buf.getvalue()
