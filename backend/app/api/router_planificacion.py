# backend/app/api/router_planificacion.py
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from typing import Optional, List
from app.core.database import supabase
from app.services.rag_orchestrator import RAGOrchestrator
from app.services.ai_service import SYSTEM_PROMPT_XLSX
from app.utils.engines import FileEngine
from app.api.generacion_utils import process_and_upload
from app.services.email_service import enviar_email
from pydantic import BaseModel
from datetime import datetime, timedelta, date
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from fastapi.responses import StreamingResponse
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from fastapi.responses import StreamingResponse


router = APIRouter()


# ── Modelos Pydantic ──────────────────────────────────────────────────────────

class ClaseWizard(BaseModel):
    numero: int
    fecha_programada: str       # "YYYY-MM-DD"
    tema_clase: str
    tipo: str                   # "clase" | "examen" | "recuperatorio"
    estado_clase: str = "programada"

class ExamenWizard(BaseModel):
    numero: int
    clases_examen: str          # "1, 2, 3"
    tiene_recuperatorio: bool = False
    clases_recup_desde: Optional[int] = None
    clases_recup_hasta: Optional[int] = None

class FeriadoWizard(BaseModel):
    fecha: str
    nombre: str

class PlanificacionWizardPayload(BaseModel):
    id_docente: str
    id_curso: str
    nombre_clase: str
    tema: str
    duracion: Optional[str] = None
    contenido_minimo: Optional[str] = None
    clases: List[ClaseWizard]
    examenes: Optional[List[ExamenWizard]] = []
    feriados_excluidos: Optional[List[FeriadoWizard]] = []

class ReplanificarClaseRequest(BaseModel):
    nueva_fecha: str            # "YYYY-MM-DD"
    motivo: str = ""
    desplazar_siguientes: bool = True  # ← clave: arrastra las clases posteriores

class EstadoClaseRequest(BaseModel):
    estado: str  # "programada" | "dictada" | "cancelada" | "reprogramada"

# ── Helpers ───────────────────────────────────────────────────────────────────

def _es_feriado(fecha_iso: str, feriados: list) -> bool:
    """Devuelve True si la fecha cae en algún feriado."""
    try:
        dt = date.fromisoformat(fecha_iso)
        for f in feriados:
            inicio = date.fromisoformat(f["fecha_inicio"][:10])
            fin    = date.fromisoformat(f["fecha_fin"][:10])
            if inicio <= dt <= fin:
                return True
    except Exception:
        pass
    return False


def _siguiente_habil(fecha_iso: str, feriados: list, dias_max: int = 60) -> str:
    """
    Dado un YYYY-MM-DD, avanza de a 1 día hasta encontrar
    una fecha que no sea feriado ni fin de semana (sáb/dom).
    Devuelve la fecha hábil como string.
    """
    dt = date.fromisoformat(fecha_iso)
    for _ in range(dias_max):
        dt += timedelta(days=1)
        iso = dt.isoformat()
        if dt.weekday() < 5 and not _es_feriado(iso, feriados):
            return iso
    return fecha_iso  # fallback: misma fecha si no encontró


def _obtener_planificacion_por_id(id_plan: str):
    """Busca la planificación por id_planificacion o por id, por compatibilidad con distintos esquemas."""
    for column in ("id_planificacion", "id"):
        try:
            res = supabase.table("planificacion").select("*").eq(column, id_plan).single().execute()
            data = getattr(res, "data", None)
            if data:
                return data
        except Exception:
            continue
    return None


def _cargar_feriados(id_docente: str) -> list:
    """Carga feriados nacionales + propios del docente."""
    try:
        nacionales = supabase.table("feriados").select("*").is_("id_docente", "null").execute().data or []
        propios    = supabase.table("feriados").select("*").eq("id_docente", id_docente).execute().data or []
        return nacionales + propios
    except Exception:
        return []


# ── Endpoints ─────────────────────────────────────────────────────────────────

# IMPORTANTE: /planificacion/wizard debe ir ANTES de /planificacion/{id}/...
# para que FastAPI no confunda "wizard" con un UUID.

@router.post("/planificacion/wizard")
async def crear_planificacion_wizard(payload: PlanificacionWizardPayload):
    """
    Recibe la planificación completa desde el wizard.
    Persiste en `planificacion` + `cronograma_clases` + `examenes_planificacion`.
    """
    try:
        curso_res = (
            supabase.table("cursos")
            .select("id_escuela, nombre_materia")
            .eq("id_curso", payload.id_curso)
            .single()
            .execute()
        )
        if not curso_res.data:
            raise HTTPException(status_code=404, detail="Curso no encontrado")

        id_escuela = curso_res.data["id_escuela"]
        fecha_principal = payload.clases[0].fecha_programada if payload.clases else None

        plan_res = supabase.table("planificacion").insert({
            "id_docente":       payload.id_docente,
            "id_escuela":       id_escuela,
            "id_curso":         payload.id_curso,
            "titulo_plan":      payload.nombre_clase or payload.tema,   # ← esto faltaba
            "nombre_clase":     payload.nombre_clase,
            "fecha":            fecha_principal,
            "duracion":         payload.duracion or "",
            "tema":             payload.tema,
            "contenido_minimo": payload.contenido_minimo or "",
            "estado":           "activa",
        }).execute()

        if not plan_res.data:
            raise HTTPException(status_code=500, detail="No se pudo crear la planificación")

        id_plan = plan_res.data[0].get("id_planificacion") or plan_res.data[0].get("id")

        # Clases individuales en cronograma_clases
        if payload.clases:
            supabase.table("cronograma_clases").insert([
                {
                    "id_planificacion": id_plan,
                    "numero":           c.numero,
                    "fecha_programada": c.fecha_programada,
                    "tema_clase":       c.tema_clase,
                    "tipo":             c.tipo,
                    "estado_clase":     c.estado_clase,
                }
                for c in payload.clases
            ]).execute()

        # Exámenes
        if payload.examenes:
            try:
                supabase.table("examenes_planificacion").insert([
                    {
                        "id_planificacion":    id_plan,
                        "numero":              ex.numero,
                        "clases_examen":       ex.clases_examen,
                        "tiene_recuperatorio": ex.tiene_recuperatorio,
                        "clases_recup_desde":  ex.clases_recup_desde,
                        "clases_recup_hasta":  ex.clases_recup_hasta,
                    }
                    for ex in payload.examenes
                ]).execute()
            except Exception as e:
                print(f"⚠️ examenes_planificacion: {e}")

        return {
            "ok": True,
            "id_planificacion": id_plan,
            "clases_creadas": len(payload.clases),
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR wizard: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/planificacion/cronograma/{id_planificacion}")
async def get_cronograma(id_planificacion: str):
    """
    Devuelve todas las clases de una planificación ordenadas por número.
    Usado por CalendarioDocente (FullCalendar).
    """
    try:
        res = (
            supabase.table("cronograma_clases")
            .select("*")
            .eq("id_planificacion", id_planificacion)
            .order("numero", desc=False)
            .execute()
        )
        return res.data or []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/planificacion/clase/{id_clase}/replanificar")
async def replanificar_clase(id_clase: str, body: ReplanificarClaseRequest):
    """
    Replanifica UNA clase individual.

    Si desplazar_siguientes=True (default):
      - Calcula el delta de días entre la fecha original y la nueva.
      - Aplica ese delta a TODAS las clases con número mayor, respetando feriados.
      - Los exámenes y recuperatorios también se desplazan porque son filas
        en cronograma_clases con tipo='examen'/'recuperatorio'.

    Si desplazar_siguientes=False:
      - Solo mueve la clase indicada.
    """
    try:
        # 1. Obtener la clase a replanificar
        clase_res = (
            supabase.table("cronograma_clases")
            .select("*")
            .eq("id", id_clase)
            .single()
            .execute()
        )
        if not clase_res.data:
            raise HTTPException(status_code=404, detail="Clase no encontrada")

        clase = clase_res.data
        id_planificacion = clase["id_planificacion"]
        numero_clase     = clase["numero"]
        fecha_original   = clase["fecha_programada"]

        # 2. Calcular delta en días
        fecha_orig_dt  = date.fromisoformat(fecha_original[:10])
        fecha_nueva_dt = date.fromisoformat(body.nueva_fecha[:10])
        delta_dias     = (fecha_nueva_dt - fecha_orig_dt).days

        # 3. Actualizar la clase indicada
        supabase.table("cronograma_clases").update({
            "fecha_programada": body.nueva_fecha,
            "estado_clase":     "reprogramada",
            "motivo_reprogramacion": body.motivo or None,
        }).eq("id", id_clase).execute()

        clases_afectadas = [{"id": id_clase, "nueva_fecha": body.nueva_fecha}]

        # 4. Si hay que desplazar las siguientes, recalcular en cascada
        if body.desplazar_siguientes and delta_dias != 0:
            # Obtener docente para cargar sus feriados
            plan_res = (
                supabase.table("planificacion")
                .select("id_docente")
                .eq("id_planificacion", id_planificacion)
                .single()
                .execute()
            )
            id_docente = (plan_res.data or {}).get("id_docente", "")
            feriados   = _cargar_feriados(id_docente)

            # Obtener todas las clases POSTERIORES a la replanificada
            siguientes_res = (
                supabase.table("cronograma_clases")
                .select("id, numero, fecha_programada, tipo")
                .eq("id_planificacion", id_planificacion)
                .gt("numero", numero_clase)
                .order("numero", desc=False)
                .execute()
            )
            siguientes = siguientes_res.data or []

            for sig in siguientes:
                try:
                    fecha_actual_dt  = date.fromisoformat(sig["fecha_programada"][:10])
                    fecha_nueva_sig  = (fecha_actual_dt + timedelta(days=delta_dias)).isoformat()

                    # Si la nueva fecha cae en feriado, avanzar al siguiente hábil
                    if _es_feriado(fecha_nueva_sig, feriados):
                        fecha_nueva_sig = _siguiente_habil(fecha_nueva_sig, feriados)

                    supabase.table("cronograma_clases").update({
                        "fecha_programada": fecha_nueva_sig,
                        # Solo marcar como reprogramada si era clase normal
                        "estado_clase": "reprogramada" if sig["tipo"] == "clase" else sig.get("estado_clase", "programada"),
                    }).eq("id", sig["id"]).execute()

                    clases_afectadas.append({"id": sig["id"], "nueva_fecha": fecha_nueva_sig})
                except Exception as e:
                    print(f"⚠️ No se pudo actualizar clase {sig['id']}: {e}")

        return {
            "ok": True,
            "clase_replanificada": id_clase,
            "fecha_original": fecha_original,
            "nueva_fecha": body.nueva_fecha,
            "delta_dias": delta_dias,
            "clases_desplazadas": len(clases_afectadas) - 1,
            "detalle": clases_afectadas,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR replanificar: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/planificacion/agenda/{id_docente}")
async def get_agenda_docente(id_docente: str):
    """Devuelve todas las planificaciones del docente con sus clases, ordenadas por fecha."""
    try:
        # Primero traemos las planificaciones
        plans_res = (
            supabase.table("planificacion")
            .select("id_planificacion, nombre_clase, tema, duracion, id_curso, id_escuela, fecha, estado")
            .eq("id_docente", id_docente)
            .order("fecha", desc=False)
            .execute()
        )
        planes = plans_res.data or []

        # Para cada planificación traemos sus clases
        resultado = []
        for plan in planes:
            clases_res = (
                supabase.table("cronograma_clases")
                .select("id, numero, fecha_programada, tema_clase, tipo, estado_clase")
                .eq("id_planificacion", plan["id_planificacion"])
                .order("numero", desc=False)
                .execute()
            )
            resultado.append({
                **plan,
                "clases": clases_res.data or [],
            })

        return {"status": "success", "agenda": resultado}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/planificacion/proximas/{id_docente}")
async def proximas_clases(id_docente: str, dias: int = 30):
    """
    Devuelve las clases individuales programadas en los próximos N días.
    Consulta cronograma_clases directamente para tener fechas actualizadas.
    """
    try:
        desde = date.today().isoformat()
        hasta = (date.today() + timedelta(days=dias)).isoformat() + "T23:59:59"

        # Traer IDs de planificaciones del docente
        plans_ids = [
            p["id_planificacion"]
            for p in (
                supabase.table("planificacion")
                .select("id_planificacion")
                .eq("id_docente", id_docente)
                .execute()
                .data or []
            )
        ]

        if not plans_ids:
            return []

        # Traer clases en el rango de fechas
        res = (
            supabase.table("cronograma_clases")
            .select("*, planificacion(nombre_clase, tema, duracion, id_curso)")
            .in_("id_planificacion", plans_ids)
            .gte("fecha_programada", desde)
            .lte("fecha_programada", hasta)
            .order("fecha_programada", desc=False)
            .execute()
        )
        return res.data or []

    except Exception as e:
        print(f"❌ proximas_clases: {e}")
        raise HTTPException(500, str(e))


@router.get("/planificacion/{id_planificacion}/detalle")
async def get_planificacion_detalle(id_planificacion: str):
    """Devuelve planificación + todas sus clases + exámenes."""
    try:
        plan = (
            supabase.table("planificacion")
            .select("*")
            .eq("id_planificacion", id_planificacion)
            .single()
            .execute()
        )
        if not plan.data:
            raise HTTPException(status_code=404, detail="Planificación no encontrada")

        clases = (
            supabase.table("cronograma_clases")
            .select("*")
            .eq("id_planificacion", id_planificacion)
            .order("numero", desc=False)
            .execute()
        )
        examenes = (
            supabase.table("examenes_planificacion")
            .select("*")
            .eq("id_planificacion", id_planificacion)
            .execute()
        )

        return {
            **plan.data,
            "clases":   clases.data or [],
            "examenes": examenes.data or [],
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Recordatorios ─────────────────────────────────────────────────────────────

@router.post("/recordatorio")
async def crear_recordatorio(
    id_planificacion: str = Form(...),
    id_docente: str = Form(...),
    minutos_antes: int = Form(...),
):
    try:
        plan = supabase.table("planificacion").select("*").eq("id_planificacion", id_planificacion).single().execute().data
        if not plan:
            raise HTTPException(404, "Planificación no encontrada")

        fecha_clase = datetime.fromisoformat(plan["fecha"])
        fecha_envio = fecha_clase - timedelta(minutes=minutos_antes)

        if fecha_envio < datetime.now():
            raise HTTPException(400, "Ese recordatorio ya pasó. Elegí menos anticipación.")

        res = supabase.table("recordatorios_clase").insert({
            "id_planificacion": id_planificacion,
            "id_docente":       id_docente,
            "minutos_antes":    minutos_antes,
            "fecha_envio":      fecha_envio.isoformat(),
        }).execute()
        return {"status": "ok", "fecha_envio": fecha_envio.isoformat(), "data": res.data[0] if res.data else {}}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Error al crear recordatorio: {str(e)}")


@router.get("/recordatorios/{id_docente}")
async def listar_recordatorios(id_docente: str):
    res = supabase.table("recordatorios_clase").select("*, planificacion(nombre_clase, fecha)").eq("id_docente", id_docente).execute()
    return res.data or []


@router.delete("/recordatorio/{id_recordatorio}")
async def eliminar_recordatorio(id_recordatorio: str):
    supabase.table("recordatorios_clase").delete().eq("id_recordatorio", id_recordatorio).execute()
    return {"status": "ok"}

@router.put("/planificacion/clase/{id_clase}/estado")
async def actualizar_estado_clase(id_clase: str, body: EstadoClaseRequest):
    """Actualiza el estado de una clase individual (dictada, cancelada, etc.)."""
    estados_validos = {"programada", "dictada", "cancelada", "reprogramada"}
    if body.estado not in estados_validos:
        raise HTTPException(status_code=400, detail=f"Estado inválido. Opciones: {estados_validos}")
    try:
        res = supabase.table("cronograma_clases") \
            .update({"estado_clase": body.estado}) \
            .eq("id", id_clase) \
            .execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Clase no encontrada")
        return {"ok": True, "id": id_clase, "estado": body.estado}
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR estado clase: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
class UnidadInput(BaseModel):
    numero: int
    nombre: str
    contenido: str                  # contenido mínimo de la unidad
    bibliografia_especifica: str = ""
 
class DistribuirPayload(BaseModel):
    id_docente: str
    nombre_asignatura: str
    contenido_minimo_general: str
    bibliografia_general: str = ""
    unidades: List[UnidadInput]
    total_clases: int               # cantidad total de clases de la asignatura
    fechas: List[str]               # lista de fechas "YYYY-MM-DD" ya calculadas
 
 
class ClaseDistribuida(BaseModel):
    numero: int
    fecha: str
    unidad: int
    tema: str
 
 
@router.post("/planificacion/distribuir")
async def distribuir_temas(payload: DistribuirPayload):
    """
    Llama a Groq para distribuir los temas de cada unidad entre
    las clases disponibles. Devuelve una lista de ClaseDistribuida.
    """
    try:
        # ── Construir prompt ──────────────────────────────────────────────────
        unidades_txt = ""
        for u in payload.unidades:
            unidades_txt += f"\n### Unidad {u.numero}: {u.nombre}\n"
            unidades_txt += f"Contenido mínimo:\n{u.contenido}\n"
            if u.bibliografia_especifica:
                unidades_txt += f"Bibliografía específica:\n{u.bibliografia_especifica}\n"
 
        fechas_txt = "\n".join(
            [f"Clase {i+1}: {f}" for i, f in enumerate(payload.fechas[:payload.total_clases])]
        )
 
        system_prompt = """
Sos un experto en planificación pedagógica universitaria.
Tu tarea es distribuir los contenidos de una asignatura entre sus clases.
Respondé EXCLUSIVAMENTE con un JSON válido, sin texto extra, sin markdown.
El JSON debe tener esta estructura exacta:
{
  "clases": [
    {
      "numero": 1,
      "fecha": "YYYY-MM-DD",
      "unidad": 1,
      "tema": "Nombre breve y preciso del tema de esta clase (máx 80 caracteres)"
    }
  ]
}
Reglas:
- Cubrí TODOS los contenidos mínimos de todas las unidades.
- Distribuí las clases proporcionalmente: más clases a unidades con más contenido.
- Cada clase tiene UN solo tema principal, concreto y específico (no genérico).
- Los temas deben seguir el orden lógico de la asignatura.
- Podés usar varias clases para un mismo subtema si es complejo.
- No inventés temas que no estén en el contenido mínimo.
- Respetá el número y fecha de cada clase tal como se indica.
"""
 
        user_prompt = f"""
Asignatura: {payload.nombre_asignatura}
 
Contenido mínimo general:
{payload.contenido_minimo_general}
 
Bibliografía general:
{payload.bibliografia_general or "No especificada"}
 
Unidades:
{unidades_txt}
 
Clases disponibles ({payload.total_clases} en total):
{fechas_txt}
 
Distribuí los temas de todas las unidades entre estas {payload.total_clases} clases.
"""
 
        # ── Llamar a Groq ─────────────────────────────────────────────────────
        from app.services.rag_orchestrator import RAGOrchestrator
        import json, re
 
        full_prompt = f"{system_prompt}\n\n{user_prompt}"
        raw = RAGOrchestrator._generate(full_prompt)
 
        # Parsear JSON
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if not match:
            raise HTTPException(500, "La IA no devolvió un JSON válido")
        data = json.loads(match.group(0))
 
        clases = data.get("clases", [])
        if not clases:
            raise HTTPException(500, "La IA no generó clases")
 
        return {"ok": True, "clases": clases}
 
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR distribuir_temas: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/planificacion/{id_plan}/exportar-word")
async def exportar_planificacion_word(id_plan: str):
    try:
        # Obtener datos de la planificación
        plan = _obtener_planificacion_por_id(id_plan)
        if not plan:
            raise HTTPException(status_code=404, detail="Planificación no encontrada")
        clases_res = supabase.table("cronograma_clases") \
            .select("*") \
            .eq("id_planificacion", id_plan) \
            .order("numero") \
            .execute()
        clases = clases_res.data or []
        # Construir el documento Word
        doc = Document()
        # Título
        titulo = doc.add_heading(plan.get("nombre_clase", "Planificación"), 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Datos generales
        doc.add_heading("Datos generales", level=1)
        tabla_meta = doc.add_table(rows=4, cols=2)
        tabla_meta.style = "Table Grid"
        celdas = [
            ("Docente",          plan.get("id_docente", "—")),
            ("Contenido mínimo", plan.get("contenido_minimo", "—")),
            ("Duración",         plan.get("duracion", "—")),
            ("Total de clases",  str(len([c for c in clases if c.get("tipo") == "clase"]))),
        ]
        for i, (k, v) in enumerate(celdas):
            tabla_meta.rows[i].cells[0].text = k
            tabla_meta.rows[i].cells[1].text = v
        doc.add_paragraph()
        # Cronograma
        doc.add_heading("Cronograma de clases", level=1)
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Table Grid"
        encabezados = ["N°", "Fecha", "Tipo", "Tema"]
        for i, enc in enumerate(encabezados):
            cell = tabla.rows[0].cells[i]
            cell.text = enc
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        tipo_labels = {"clase": "Clase", "examen": "Examen", "recuperatorio": "Recuperatorio"}
        for c in clases:
            row = tabla.add_row()
            row.cells[0].text = str(c.get("numero", ""))
            row.cells[1].text = c.get("fecha_programada", "")
            row.cells[2].text = tipo_labels.get(c.get("tipo", "clase"), c.get("tipo", ""))
            row.cells[3].text = c.get("tema_clase", "")
        # Guardar en buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        nombre_archivo = f"Planificacion_{plan.get('nombre_clase', id_plan).replace(' ', '_')}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={nombre_archivo}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR exportar word: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ── Listar planificaciones del docente ────────────────────────────────────────
@router.get("/planificacion/lista/{id_docente}")
async def listar_planificaciones(id_docente: str):
    try:
        res = supabase.table("planificacion") \
            .select("id, nombre_clase, duracion, contenido_minimo, created_at") \
            .eq("id_docente", id_docente) \
            .order("created_at", desc=True) \
            .execute()
        plans = res.data or []
        # Agregar total_clases de cronograma_clases
        for p in plans:
            try:
                clases_res = supabase.table("cronograma_clases") \
                    .select("id", count="exact") \
                    .eq("id_planificacion", p["id"]) \
                    .execute()
                p["total_clases"] = clases_res.count or 0
            except Exception:
                p["total_clases"] = 0
        return plans
    except Exception as e:
        print(f"❌ ERROR listar planificaciones: {e}")
        raise HTTPException(status_code=500, detail=str(e))
# ── Exportar planificación a Word (.docx) ─────────────────────────────────────
@router.get("/planificacion/{id_plan}/exportar-word")
async def exportar_planificacion_word(id_plan: str):
    try:
        plan = _obtener_planificacion_por_id(id_plan)
        if not plan:
            raise HTTPException(status_code=404, detail="Planificación no encontrada")
        clases_res = supabase.table("cronograma_clases") \
            .select("*").eq("id_planificacion", id_plan).order("numero").execute()
        clases = clases_res.data or []
        doc = Document()
        # Título
        titulo = doc.add_heading(plan.get("nombre_clase", "Planificación"), 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Datos generales
        doc.add_heading("Datos generales", level=1)
        tabla_meta = doc.add_table(rows=3, cols=2)
        tabla_meta.style = "Table Grid"
        meta = [
            ("Contenido mínimo", plan.get("contenido_minimo", "—")),
            ("Duración de clase",  plan.get("duracion", "—")),
            ("Total de clases",    str(len([c for c in clases if c.get("tipo") == "clase"]))),
        ]
        for i, (k, v) in enumerate(meta):
            tabla_meta.rows[i].cells[0].text = k
            tabla_meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tabla_meta.rows[i].cells[1].text = v or "—"
        doc.add_paragraph()
        # Cronograma
        doc.add_heading("Cronograma", level=1)
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Table Grid"
        for i, enc in enumerate(["N°", "Fecha", "Tipo", "Tema"]):
            cell = tabla.rows[0].cells[i]
            cell.text = enc
            cell.paragraphs[0].runs[0].bold = True
        tipo_labels = {"clase": "Clase", "examen": "Examen", "recuperatorio": "Recuperatorio"}
        for c in clases:
            row = tabla.add_row()
            row.cells[0].text = str(c.get("numero", ""))
            row.cells[1].text = c.get("fecha_programada", "")
            row.cells[2].text = tipo_labels.get(c.get("tipo", "clase"), c.get("tipo", ""))
            row.cells[3].text = c.get("tema_clase", "")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        nombre = f"Planificacion_{plan.get('nombre_clase', id_plan).replace(' ', '_')}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={nombre}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR exportar word: {e}")
        raise HTTPException(status_code=500, detail=str(e))
# ── Exportar planificación a PDF ──────────────────────────────────────────────
@router.get("/planificacion/{id_plan}/exportar-pdf")
async def exportar_planificacion_pdf(id_plan: str):
    try:
        plan = _obtener_planificacion_por_id(id_plan)
        if not plan:
            raise HTTPException(status_code=404, detail="Planificación no encontrada")
        clases_res = supabase.table("cronograma_clases") \
            .select("*").eq("id_planificacion", id_plan).order("numero").execute()
        clases = clases_res.data or []
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=50, bottomMargin=40)
        styles = getSampleStyleSheet()
        elementos = []
        # Título
        elementos.append(Paragraph(plan.get("nombre_clase", "Planificación"), styles["Title"]))
        elementos.append(Spacer(1, 12))
        # Datos generales
        elementos.append(Paragraph("Datos generales", styles["Heading2"]))
        meta_data = [
            ["Contenido mínimo", plan.get("contenido_minimo", "—") or "—"],
            ["Duración de clase",  plan.get("duracion", "—") or "—"],
            ["Total de clases",    str(len([c for c in clases if c.get("tipo") == "clase"]))],
        ]
        t_meta = Table(meta_data, colWidths=[140, 360])
        t_meta.setStyle(TableStyle([
            ("BACKGROUND",  (0, 0), (0, -1), colors.HexColor("#e0f2fe")),
            ("FONTNAME",    (0, 0), (0, -1), "Helvetica-Bold"),
            ("GRID",        (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTSIZE",    (0, 0), (-1, -1), 9),
            ("VALIGN",      (0, 0), (-1, -1), "TOP"),
            ("PADDING",     (0, 0), (-1, -1), 6),
        ]))
        elementos.append(t_meta)
        elementos.append(Spacer(1, 16))
        # Cronograma
        elementos.append(Paragraph("Cronograma de clases", styles["Heading2"]))
        tipo_labels = {"clase": "Clase", "examen": "Examen", "recuperatorio": "Recuperatorio"}
        tipo_colores = {"clase": "#dbeafe", "examen": "#fef3c7", "recuperatorio": "#dcfce7"}
        tabla_data = [["N°", "Fecha", "Tipo", "Tema"]]
        for c in clases:
            tabla_data.append([
                str(c.get("numero", "")),
                c.get("fecha_programada", ""),
                tipo_labels.get(c.get("tipo", ""), c.get("tipo", "")),
                Paragraph(c.get("tema_clase", ""), styles["Normal"]),
            ])
        t = Table(tabla_data, colWidths=[25, 75, 80, 320])
        style_cmds = [
            ("BACKGROUND",  (0, 0), (-1, 0),  colors.HexColor("#1e3a8a")),
            ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, -1), 8),
            ("GRID",        (0, 0), (-1, -1), 0.4, colors.grey),
            ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING",     (0, 0), (-1, -1), 5),
        ]
        for i, c in enumerate(clases, start=1):
            bg = tipo_colores.get(c.get("tipo", "clase"), "#ffffff")
            style_cmds.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor(bg)))
        t.setStyle(TableStyle(style_cmds))
        elementos.append(t)
        doc.build(elementos)
        buf.seek(0)
        nombre = f"Planificacion_{plan.get('nombre_clase', id_plan).replace(' ', '_')}.pdf"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={nombre}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ERROR exportar pdf: {e}")
        raise HTTPException(status_code=500, detail=str(e))