# backend/app/api/contenido_helpers.py
"""
Funciones auxiliares del módulo de Generación de Contenido: obtención de
datos de escuela/materia, armado de contexto bibliográfico, y encabezado
institucional de los documentos Word.
Extraído sin cambios de lógica desde router_contenido.py.
"""
from typing import Optional
from datetime import datetime
from app.core.database import supabase


def _datos_escuela_materia(id_escuela: Optional[str], id_curso: Optional[str]):
    """
    Devuelve (nombre_escuela, nombre_materia, division, contenido_minimo, bibliografia).
    Consulta Supabase solo si se proveen los IDs.
    """
    nombre_escuela   = "Institución Educativa"
    nombre_materia   = "Materia General"
    division         = "-"
    contenido_minimo = ""
    bibliografia     = []

    if id_escuela:
        try:
            res = supabase.table("escuelas").select("nombre_escuela") \
                .eq("id_escuela", id_escuela).single().execute()
            nombre_escuela = (res.data or {}).get("nombre_escuela", nombre_escuela)
        except Exception as e:
            print(f"⚠️ No pude obtener escuela: {e}")

    if id_curso:
        try:
            res = supabase.table("cursos") \
                .select("nombre_materia,division,contenido_minimo,bibliografia") \
                .eq("id_curso", id_curso).single().execute()
            if res.data:
                nombre_materia   = res.data.get("nombre_materia",   nombre_materia)
                division         = res.data.get("division",         division)
                contenido_minimo = res.data.get("contenido_minimo", "") or ""
                bibliografia     = res.data.get("bibliografia",     []) or []
        except Exception as e:
            print(f"⚠️ No pude obtener curso: {e}")

    return nombre_escuela, nombre_materia, division, contenido_minimo, bibliografia


def _contexto_biblio(contenido_minimo: str, bibliografia: list) -> str:
    """Arma texto de contexto cuando no hay PDF."""
    partes = []
    if contenido_minimo:
        partes.append(f"Contenido mínimo de la materia:\n{contenido_minimo}")
    if bibliografia:
        items = [b if isinstance(b, str) else str(b) for b in bibliografia]
        partes.append("Bibliografía de la materia:\n" + "\n".join(f"- {i}" for i in items))
    return "\n\n".join(partes) if partes else \
        "Sin contenido específico. Generá basándote en el tema indicado."


def _encabezado_documento(doc, tipo_doc: str, nombre_materia: str, tema: str,
                           nombre_escuela: str, fecha_str: Optional[str] = None):
    """
    Agrega el encabezado institucional al documento Word.
    Formato (imagen de referencia):
        [Centrado, negrita, 18pt] {tipo_doc} de {nombre_materia}
        [Centrado, negrita, 14pt] Tema: {tema}
        <línea en blanco>
        Escuela: {nombre_escuela}
        Nombre del alumno: ______________________________
        Fecha: {fecha}
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # If fecha_str is None -> use current date; if it's an empty string -> keep it empty
    if fecha_str is None:
        fecha = datetime.now().strftime("%d/%m/%Y")
    else:
        fecha = fecha_str

    # Título principal centrado
    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo_p.add_run(f"{tipo_doc} de {nombre_materia}")
    run_titulo.bold = True
    run_titulo.font.size = Pt(18)

    # Subtítulo Tema centrado
    tema_p = doc.add_paragraph()
    tema_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tema = tema_p.add_run(f"Tema: {tema}")
    run_tema.bold = True
    run_tema.font.size = Pt(14)

    doc.add_paragraph()  # espacio en blanco

    # Escuela
    p_esc = doc.add_paragraph()
    p_esc.add_run("Escuela: ").bold = True
    p_esc.add_run(nombre_escuela)

    # Nombre del alumno (siempre en blanco para que el alumno complete a mano)
    p_alumno = doc.add_paragraph()
    p_alumno.add_run("Nombre del alumno: ").bold = True
    p_alumno.add_run("______________________________")

    # Fecha
    p_fecha = doc.add_paragraph()
    p_fecha.add_run("Fecha: ").bold = True
    p_fecha.add_run(fecha)

    doc.add_paragraph()  # espacio antes del contenido
