# backend/app/services/planning_service.py
import holidays
from datetime import datetime, timedelta, time
from typing import List, Dict, Optional
import uuid
from app.core.database import supabase
from app.models.schemas import ClaseSchema, ExamenSchema, RecuperatorioSchema

class PlanningService:
    def __init__(self, country='AR'):  # Argentina por defecto
        self.holidays = holidays.CountryHoliday(country)

    def is_holiday(self, date: datetime.date) -> bool:
        return date in self.holidays

    def generate_schedule(self, request):
        """
        Genera el cronograma académico inteligente.
        """
        # Parsear horarios
        schedule_slots = self.parse_schedules(request.horarios)

        # Obtener fechas disponibles
        available_dates = self.get_available_dates(
            schedule_slots, 
            request.cantidad_clases + request.cantidad_examenes + request.cantidad_recuperatorios
        )

        # Generar clases
        clases = []
        examenes = []
        recuperatorios = []

        date_index = 0

        # Clases
        for i in range(request.cantidad_clases):
            # Asignar bibliografía específica: dividir la bibliografía en capítulos
            bibliografia_parts = request.bibliografia.split('\n') if request.bibliografia else []
            bib_especifica = bibliografia_parts[i % len(bibliografia_parts)] if bibliografia_parts else request.bibliografia

            clase = ClaseSchema(
                id_clase=str(uuid.uuid4()),
                orden=i+1,
                fecha_programada=available_dates[date_index],
                tema_clase=f"Clase {i+1}",
                actividades_previstas=request.contenido_minimo,
                bibliografia_especifica=bib_especifica,
                estado_clase="pendiente"
            )
            clases.append(clase)
            date_index += 1

        # Exámenes
        for i in range(request.cantidad_examenes):
            # Asignar clases dinámicamente: examen 1 evalúa clases 1-3, examen 2 clases 4-6, etc.
            start_clase = i * 3
            end_clase = min((i+1) * 3, len(clases))
            clases_ids = [c.id_clase for c in clases[start_clase:end_clase]]

            examen = ExamenSchema(
                id_examen=str(uuid.uuid4()),
                id_curso=request.id_curso,
                nombre=f"Examen {i+1}",
                fecha=available_dates[date_index],
                clases_ids=clases_ids,
                tiene_recuperatorio=i < request.cantidad_recuperatorios
            )
            examenes.append(examen)
            date_index += 1

        # Recuperatorios
        for i in range(request.cantidad_recuperatorios):
            # Recuperatorio evalúa todas las clases del bloque correspondiente
            examen = examenes[i]
            recuperatorio = RecuperatorioSchema(
                id_recuperatorio=str(uuid.uuid4()),
                id_examen=examen.id_examen,
                fecha=available_dates[date_index],
                clases_ids=examen.clases_ids  # Mismas clases que el examen
            )
            recuperatorios.append(recuperatorio)
            date_index += 1

        # Resolver conflictos y sugerencias
        self.resolve_conflicts(clases, examenes, recuperatorios, schedule_slots)

        return {
            'clases': clases,
            'examenes': examenes,
            'recuperatorios': recuperatorios
        }

    def parse_schedules(self, horarios: List[str]) -> List[Dict]:
        """
        Parsea los horarios en slots disponibles.
        """
        slots = []
        for h in horarios:
            # Ej: "lunes 10:00-11:00"
            parts = h.split()
            day = parts[0].lower()
            time_range = parts[1].split('-')
            start_time = time.fromisoformat(time_range[0])
            end_time = time.fromisoformat(time_range[1])
            slots.append({
                'day': day,
                'start': start_time,
                'end': end_time
            })
        return slots

    def get_available_dates(self, slots: List[Dict], total_events: int) -> List[datetime]:
        """
        Obtiene fechas disponibles evitando feriados y colisiones.
        """
        dates = []
        current_date = datetime.now().date()
        day_map = {
            'lunes': 0, 'martes': 1, 'miercoles': 2, 'jueves': 3, 'viernes': 4, 'sabado': 5, 'domingo': 6
        }

        while len(dates) < total_events:
            for slot in slots:
                if current_date.weekday() == day_map.get(slot['day'], -1):
                    if not self.is_holiday(current_date):
                        dt = datetime.combine(current_date, slot['start'])
                        dates.append(dt)
                        if len(dates) >= total_events:
                            break
            current_date += timedelta(days=1)

        return dates

    def resolve_conflicts(self, clases, examenes, recuperatorios, slots):
        """
        Resuelve conflictos y aplica reglas de reacomodamiento.
        """
        # Lógica básica: si hay recuperatorio, reservar espacio y desplazar clases posteriores
        for rec in recuperatorios:
            examen = next(e for e in examenes if e.id_examen == rec.id_examen)
            if rec.fecha <= examen.fecha:
                # Sugerir mover recuperatorio
                rec.fecha = self.find_next_available(examen.fecha + timedelta(days=1), slots)

        # Para conflictos generales, sugerir mover a fecha más cercana
        all_events = clases + examenes + recuperatorios
        sorted_events = sorted(all_events, key=lambda x: x.fecha_programada if hasattr(x, 'fecha_programada') else x.fecha)

        for i in range(1, len(sorted_events)):
            if sorted_events[i].fecha <= sorted_events[i-1].fecha:
                # Conflicto, mover al siguiente disponible
                next_date = self.find_next_available(sorted_events[i-1].fecha + timedelta(hours=1), slots)
                if hasattr(sorted_events[i], 'fecha_programada'):
                    sorted_events[i].fecha_programada = next_date
                else:
                    sorted_events[i].fecha = next_date

    def find_next_available(self, start_date: datetime, slots: List[Dict]) -> datetime:
        """
        Encuentra la siguiente fecha disponible.
        """
        current = start_date
        day_map = {'lunes': 0, 'martes': 1, 'miercoles': 2, 'jueves': 3, 'viernes': 4, 'sabado': 5, 'domingo': 6}
        while True:
            for slot in slots:
                if current.weekday() == day_map.get(slot['day'], -1) and not self.is_holiday(current.date()):
                    return current.replace(hour=slot['start'].hour, minute=slot['start'].minute)
            current += timedelta(days=1)

    def reprogramar_clase(self, id_clase: str, nueva_fecha: datetime, desplazar_subsiguientes: bool, id_curso: str):
        """
        Reprograma una clase y opcionalmente desplaza las subsiguientes.
        """
        # Obtener todas las clases del curso ordenadas por orden
        clases_data = supabase.table('clases').select('*').eq('id_curso', id_curso).order('orden').execute()
        clases = [ClaseSchema(**c) for c in clases_data.data]

        # Encontrar la clase a reprogramar
        clase_idx = next((i for i, c in enumerate(clases) if c.id_clase == id_clase), None)
        if clase_idx is None:
            raise ValueError("Clase no encontrada")

        # Calcular el delta de tiempo
        delta = nueva_fecha - clases[clase_idx].fecha_programada

        # Actualizar la clase
        supabase.table('clases').update({
            'fecha_programada': nueva_fecha.isoformat(),
            'estado_clase': 'reprogramada'
        }).eq('id_clase', id_clase).execute()

        if desplazar_subsiguientes:
            # Desplazar todas las subsiguientes
            for i in range(clase_idx + 1, len(clases)):
                nueva_fecha_sub = clases[i].fecha_programada + delta
                supabase.table('clases').update({
                    'fecha_programada': nueva_fecha_sub.isoformat(),
                    'estado_clase': 'reprogramada'
                }).eq('id_clase', clases[i].id_clase).execute()

        return {"message": "Clase reprogramada exitosamente"}
        """
        Exporta la planificación a Word.
        """
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading('Planificación Académica', 0)

        # Tabla de clases
        doc.add_heading('Cronograma de Clases', level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Clase'
        hdr_cells[1].text = 'Fecha'
        hdr_cells[2].text = 'Tema'
        hdr_cells[3].text = 'Bibliografía'

        for clase in plan_data['clases']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(clase.get('orden', clase.get('id_clase', '')))
            row_cells[1].text = clase['fecha_programada'].strftime('%Y-%m-%d')
            row_cells[2].text = clase['tema_clase']
            row_cells[3].text = clase.get('bibliografia_especifica', bibliografia)

        # Checklist de evaluaciones
        doc.add_heading('Evaluaciones', level=1)
        for examen in plan_data['examenes']:
            doc.add_paragraph(f"Examen: {examen.nombre} - Fecha: {examen.fecha.strftime('%Y-%m-%d')}")
            doc.add_paragraph("Clases evaluadas: " + ', '.join(examen.clases_ids))
            if examen.tiene_recuperatorio:
                rec = next(r for r in plan_data['recuperatorios'] if r.id_examen == examen.id_examen)
                doc.add_paragraph(f"Recuperatorio: Fecha {rec.fecha.strftime('%Y-%m-%d')} - Clases: {', '.join(rec.clases_ids)}")

        return doc